import streamlit as st
import logging
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from fpdf import FPDF
from bs4 import BeautifulSoup
import io

logger = logging.getLogger(__name__)

class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

def process_mermaid_diagram(content: str) -> tuple[str, bytes]:
    """Extract and convert Mermaid diagram to image"""
    if '```mermaid' in content:
        start = content.find('```mermaid') + 10
        end = content.find('```', start)
        if end != -1:
            diagram_code = content[start:end].strip()
            # Convert diagram to image
            try:
                from utils.mermaid_helper import render_mermaid_to_image
                diagram_image = render_mermaid_to_image(diagram_code)
                # Replace Mermaid code with placeholder
                content = content[:content.find('```mermaid')] + \
                         '[DIAGRAM_PLACEHOLDER]' + \
                         content[end + 3:]
                return content, diagram_image
            except Exception as e:
                logger.error(f"Error rendering Mermaid diagram: {str(e)}")
    return content, None

def extract_table_data(html_table: str) -> list:
    """Extract data from HTML table string"""
    try:
        soup = BeautifulSoup(html_table, 'html.parser')
        rows = []
        
        # Get headers
        headers = []
        for th in soup.find_all('th'):
            headers.append(th.text.strip())
        if headers:
            rows.append(headers)
            
        # Get data rows
        for tr in soup.find_all('tr')[1:]:  # Skip header row
            row = []
            for td in tr.find_all('td'):
                row.append(td.text.strip())
            if row:
                rows.append(row)
                
        return rows
    except:
        return None

def create_pdf_table(pdf: PDF, rows: list):
    """Create PDF table with proper column widths"""
    if not rows:
        return
        
    # Calculate optimal column widths
    n_cols = len(rows[0])
    page_width = pdf.get_page_width() - 40  # Account for margins
    col_width = page_width / n_cols
    
    # Table header
    pdf.set_font('Arial', 'B', 10)
    for header in rows[0]:
        pdf.cell(col_width, 7, header.strip(), border=1)
    pdf.ln()
    
    # Table data
    pdf.set_font('Arial', '', 10)
    for row in rows[1:]:
        for cell in row:
            pdf.cell(col_width, 6, cell.strip(), border=1)
        pdf.ln()
    
    pdf.ln(5)

def process_text_content(pdf: PDF, content: str):
    """Process text content with tables and formatting"""
    paragraphs = content.split('\n')
    
    for para in paragraphs:
        if not para.strip():
            pdf.ln(5)
            continue
            
        if para.strip().startswith('<table'):
            try:
                rows = extract_table_data(para)
                if rows:
                    create_pdf_table(pdf, rows)
                    pdf.ln(5)
            except:
                # If table processing fails, treat as normal text
                pdf.multi_cell(0, 5, para.strip())
        else:
            # Handle italic text
            parts = para.split('*')
            for i, part in enumerate(parts):
                if part.strip():
                    pdf.set_font('Arial', 'I' if i % 2 else '', 11)
                    pdf.multi_cell(0, 5, part.strip())
            pdf.ln(3)

def create_pdf(generated_sections: dict) -> PDF:
    """Create PDF document with proper formatting"""
    try:
        # Initialize PDF
        pdf = PDF()
        pdf.add_page()
        
        # Title page
        pdf.set_font('Arial', 'B', 24)
        pdf.cell(0, 20, 'Study Protocol', align='C', ln=True)
        
        # Add date
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, f'Generated: {time.strftime("%B %d, %Y")}', align='C', ln=True)
        
        # Table of contents
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, 'Table of Contents', ln=True)
        
        # Track page numbers for TOC
        section_pages = {}
        current_page = 3  # Start after TOC
        
        # Add TOC entries
        pdf.set_font('Arial', '', 12)
        for section in generated_sections:
            section_title = section.replace('_', ' ').title()
            section_pages[section] = current_page
            dots = '.' * (50 - len(section_title))
            pdf.cell(0, 8, f'{section_title} {dots} {current_page}', ln=True)
            current_page += 1
        
        # Content pages
        for section, content in generated_sections.items():
            pdf.add_page()
            
            # Section heading
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, section.replace('_', ' ').title(), ln=True)
            pdf.ln(5)
            
            # Content processing
            content, diagram = process_mermaid_diagram(content)
            
            # Add diagram if present
            if diagram:
                pdf.image(io.BytesIO(diagram), x=10, w=190)
                pdf.ln(5)
            
            # Process remaining content
            parts = content.split('<table')
            
            # Add first part (text before first table)
            if parts[0].strip():
                process_text_content(pdf, parts[0])
            
            # Process tables and remaining text
            for i, part in enumerate(parts[1:], 1):
                table_end = part.find('</table>')
                if table_end != -1:
                    # Extract and convert table
                    table_html = '<table' + part[:table_end + 8]
                    table_data = extract_table_data(table_html)
                    if table_data:
                        create_pdf_table(pdf, table_data)
                    
                    # Process remaining text
                    remaining_text = part[table_end + 8:].strip()
                    if remaining_text:
                        process_text_content(pdf, remaining_text)
        
        return pdf
        
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}")
        raise Exception(f"Failed to create PDF document: {str(e)}")

def add_text_with_formatting(doc: Document, text: str):
    """Add text to document with proper formatting"""
    for para in text.split('\n'):
        if para.strip():
            p = doc.add_paragraph()
            parts = para.split('*')
            for i, part in enumerate(parts):
                if part.strip():
                    run = p.add_run(part.strip())
                    if i % 2:  # Odd indices are italic
                        run.italic = True

def add_table_to_doc(doc: Document, table_data: list):
    """Add table to document with proper formatting"""
    if not table_data:
        return
        
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    # Add data
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            # Add cell content with bold headers
            text = cell.strip()
            cell = table.cell(i, j)
            if i == 0:  # Header row
                cell.paragraphs[0].runs[0].bold = True
            cell.text = text
            
    doc.add_paragraph()  # Add spacing after table

def create_docx(generated_sections: dict) -> Document:
    """Create DOCX document with proper formatting"""
    doc = Document()
    
    # Add title
    doc.add_heading('Study Protocol', 0)
    
    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    for section in generated_sections:
        doc.add_paragraph(
            section.replace('_', ' ').title(),
            style='TOC 1'
        )
    doc.add_page_break()
    
    # Process sections
    for section, content in generated_sections.items():
        # Add section heading
        doc.add_heading(section.replace('_', ' ').title(), level=1)
        
        # Process content
        content, diagram = process_mermaid_diagram(content)
        
        # Add diagram if present
        if diagram:
            doc.add_picture(io.BytesIO(diagram))
            doc.add_paragraph()  # Add spacing
            
        # Process tables and text
        parts = content.split('<table')
        
        # Add first part (text before first table)
        if parts[0].strip():
            add_text_with_formatting(doc, parts[0])
            
        # Process tables
        for i, part in enumerate(parts[1:], 1):
            table_end = part.find('</table>')
            if table_end != -1:
                # Extract and convert table
                table_html = '<table' + part[:table_end + 8]
                table_data = extract_table_data(table_html)
                if table_data:
                    add_table_to_doc(doc, table_data)
                
                # Add remaining text
                remaining_text = part[table_end + 8:].strip()
                if remaining_text:
                    add_text_with_formatting(doc, remaining_text)
                    
    return doc

def check_connection():
    """Check if application can establish necessary connections"""
    try:
        from utils.template_section_generator import TemplateSectionGenerator
        generator = TemplateSectionGenerator()
        return True
    except Exception as e:
        logger.error(f"Connection check failed: {str(e)}")
        return False

def render_navigator():
    """Render the section navigator with improved generation tracking"""
    try:
        # Initialize session state for progress tracking
        if 'generation_in_progress' not in st.session_state:
            st.session_state.generation_in_progress = False
        if 'section_status' not in st.session_state:
            st.session_state.section_status = {}
            
        # Check connection
        if not check_connection():
            st.sidebar.error("⚠️ Connection issues detected. Please refresh the page.")
            return
            
        st.sidebar.markdown("## Protocol Development")
        
        # Study type detection and analysis
        if synopsis_content := st.session_state.get('synopsis_content'):
            from utils.synopsis_validator import SynopsisValidator
            validator = SynopsisValidator()
            validation_result = validator.validate_synopsis(synopsis_content)
            
            if validation_result:
                if study_type := validation_result.get('study_type'):
                    st.sidebar.success(f"📋 Study Type: {study_type.replace('_', ' ').title()}")
                    if study_type in ['phase1', 'phase2', 'phase3', 'phase4']:
                        st.sidebar.info(f"🔬 Clinical Trial Phase: {study_type[-1]}")
                else:
                    st.sidebar.warning("⚠️ Could not detect study type")
                    
                if therapeutic_area := validation_result.get('therapeutic_area'):
                    st.sidebar.info(f"🏥 Therapeutic Area: {therapeutic_area.replace('_', ' ').title()}")
                    
            # Section Navigation
            study_type = st.session_state.get('study_type')
            if study_type:
                st.sidebar.markdown("### 📑 Protocol Sections")
                
                from config.study_type_definitions import COMPREHENSIVE_STUDY_CONFIGS
                study_config = COMPREHENSIVE_STUDY_CONFIGS.get(study_type, {})
                sections = study_config.get('required_sections', [])
                
                # Progress tracking
                completed = sum(1 for section in sections 
                              if st.session_state.section_status.get(section, {}).get('status') == 'completed')
                total = len(sections)
                progress = completed / total if total > 0 else 0
                
                # Progress section
                with st.sidebar.container():
                    st.markdown('<div class="progress-section">', unsafe_allow_html=True)
                    st.progress(progress, text=f"Progress: {completed}/{total} sections")
                    st.markdown('</div>', unsafe_allow_html=True)
                
                # Generate Protocol button
                if not st.session_state.generation_in_progress:
                    if st.sidebar.button(
                        "Generate Complete Protocol",
                        key="generate_btn",
                        use_container_width=True
                    ):
                        st.session_state.generation_in_progress = True
                        progress_bar = st.sidebar.progress(0)
                        status_area = st.sidebar.empty()
                        
                        try:
                            from utils.template_section_generator import TemplateSectionGenerator
                            generator = TemplateSectionGenerator()
                            sections = study_config.get('required_sections', [])
                            
                            generated_sections = {}
                            for idx, section in enumerate(sections):
                                progress = (idx + 1) / len(sections)
                                progress_bar.progress(progress)
                                status_area.info(f"Generating {section.replace('_', ' ').title()}...")
                                
                                try:
                                    content = generator.generate_section(
                                        section_name=section,
                                        synopsis_content=synopsis_content,
                                        study_type=study_type
                                    )
                                    if content:
                                        generated_sections[section] = content
                                        st.session_state.section_status[section] = {
                                            'status': 'completed',
                                            'timestamp': time.strftime('%H:%M:%S')
                                        }
                                except Exception as e:
                                    logger.error(f"Error generating {section}: {str(e)}")
                                    st.session_state.section_status[section] = {
                                        'status': 'failed',
                                        'error': str(e)
                                    }
                                    continue
                            
                            st.session_state.generated_sections = generated_sections
                            st.session_state.generation_in_progress = False
                            st.rerun()
                            
                        except Exception as e:
                            logger.error(f"Error in protocol generation: {str(e)}")
                            st.session_state.generation_in_progress = False
                            st.sidebar.error(f"Error: {str(e)}")
                
                # Section navigation and download options
                if generated_sections := st.session_state.get('generated_sections'):
                    # Display sections
                    st.sidebar.markdown("### 📑 Generated Sections")
                    for section in sections:
                        if content := generated_sections.get(section):
                            with st.sidebar.expander(f"📝 {section.replace('_', ' ').title()}"):
                                st.text_area(
                                    "Content",
                                    value=content,
                                    height=150,
                                    disabled=True,
                                    key=f"nav_{section}"
                                )
                    
                    # Download options
                    st.sidebar.markdown("### 📥 Download Protocol")
                    
                    try:
                        # Create temporary directory if it doesn't exist
                        if not os.path.exists('temp'):
                            os.makedirs('temp')
                        
                        # Generate DOCX first
                        doc = create_docx(generated_sections)
                        temp_docx = 'temp/protocol.docx'
                        doc.save(temp_docx)
                        
                        # Add DOCX download button
                        with open(temp_docx, 'rb') as f:
                            docx_data = f.read()
                            st.sidebar.download_button(
                                label="📄 Download DOCX",
                                data=docx_data,
                                file_name="protocol.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        
                        # Generate and download PDF
                        try:
                            pdf = create_pdf(generated_sections)
                            temp_pdf = 'temp/protocol.pdf'
                            pdf.output(temp_pdf)
                            
                            with open(temp_pdf, 'rb') as f:
                                pdf_data = f.read()
                                st.sidebar.download_button(
                                    label="📥 Download PDF",
                                    data=pdf_data,
                                    file_name="protocol.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                            
                            # Cleanup PDF file
                            os.remove(temp_pdf)
                            
                        except Exception as e:
                            logger.error(f"Error creating PDF: {str(e)}")
                            st.sidebar.warning("⚠️ PDF generation failed. Please try DOCX format.")
                        
                        # Cleanup DOCX file
                        os.remove(temp_docx)
                        
                    except Exception as e:
                        logger.error(f"Error preparing downloads: {str(e)}")
                        st.sidebar.error("Error preparing downloads. Please try again.")
                        
    except Exception as e:
        logger.error(f"Error in navigator: {str(e)}")
        st.error(f"An error occurred: {str(e)}")
