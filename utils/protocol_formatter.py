import os
import logging
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

class ProtocolFormatter:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.setup_custom_styles()

        # Define section order for different study types
        self.section_orders = {
            "phase1": [
                "title",
                "background",
                "objectives",
                "study_design",
                "population",
                "procedures",
                "statistical",
                "safety"
            ],
            "slr": [
                "title",
                "background",
                "objectives",
                "study_design",
                "search_strategy",
                "selection_criteria",
                "data_extraction",
                "quality_assessment",
                "statistical"
            ],
            "meta": [
                "title",
                "background",
                "objectives",
                "study_design",
                "search_strategy",
                "selection_criteria",
                "data_extraction",
                "quality_assessment",
                "statistical"
            ],
            "rwe": [
                "title",
                "background",
                "objectives",
                "study_design",
                "data_sources",
                "population",
                "variables",
                "statistical",
                "limitations"
            ]
        }

        # Define section titles mapping
        self.section_titles = {
            "background": "Background",
            "objectives": "Objectives",
            "study_design": "Study Design",
            "population": "Study Population",
            "procedures": "Study Procedures",
            "statistical": "Statistical Analysis",
            "safety": "Safety",
            "search_strategy": "Search Strategy",
            "selection_criteria": "Selection Criteria",
            "data_extraction": "Data Extraction",
            "quality_assessment": "Quality Assessment",
            "data_sources": "Data Sources",
            "variables": "Study Variables",
            "limitations": "Study Limitations"
        }

    def setup_document(self):
        """Initialize document settings"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def setup_custom_styles(self):
        '''Set up custom document styles'''
        styles = {
            'CustomTitle': {'name': 'Protocol Title', 'size': 24, 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
            'CustomHeading1': {'name': 'Section Title', 'size': 16, 'bold': True},
            'CustomHeading2': {'name': 'Subsection Title', 'size': 14, 'bold': True},
            'CustomHeading3': {'name': 'Sub-subsection Title', 'size': 12, 'bold': True},
            'CustomBody': {'name': 'Body Text', 'size': 11},
            'CustomTableText': {'name': 'Table Text', 'size': 10},
            'CustomTableHeader': {'name': 'Table Header', 'size': 10, 'bold': True},
            'CustomCaption': {'name': 'Figure Caption', 'size': 10, 'italic': True}
        }

        for style_name, props in styles.items():
            try:
                if style_name not in self.doc.styles:
                    style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                else:
                    style = self.doc.styles[style_name]
                
                font = style.font
                font.name = 'Arial'
                font.size = Pt(props['size'])
                font.bold = props.get('bold', False)
                font.italic = props.get('italic', False)
                
                para_format = style.paragraph_format
                para_format.space_before = Pt(6)
                para_format.space_after = Pt(6)
                para_format.line_spacing = 1.15
                
                if props.get('align'):
                    para_format.alignment = props['align']
                    
            except Exception as e:
                logger.error(f"Error setting up style {style_name}: {str(e)}")

    def _clean_section_content(self, content):
        """Clean up section content"""
        if not content:
            return ""

        # Remove duplicate section titles
        lines = content.split('\n')
        cleaned_lines = []
        seen_titles = set()

        for line in lines:
            # Check if line is a title
            is_title = bool(re.match(r'^#+\s+|^[A-Z\s]+$', line.strip()))
            if is_title:
                clean_title = re.sub(r'^#+\s+', '', line).strip().lower()
                if clean_title in seen_titles:
                    continue
                seen_titles.add(clean_title)
            cleaned_lines.append(line)

        content = '\n'.join(cleaned_lines)

        # Fix list formatting
        content = re.sub(r'^\s*•\s+', '• ', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*-\s+', '• ', content, flags=re.MULTILINE)

        # Fix spacing
        content = re.sub(r'\n{3,}', '\n\n', content)
        content = content.strip()

        return content

    def _format_section(self, section_name, content, section_number):
        """Format a single section"""
        if not content:
            return

        # Add section heading
        title = self.section_titles.get(section_name, section_name.replace('_', ' ').title())
        heading = self.doc.add_paragraph(f"{section_number}. {title}", style='CustomHeading1')

        # Process content
        paragraphs = content.split('\n\n')
        subsection_number = 1

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check for subsection headings
            if para.startswith('#'):
                level = len(re.match(r'^#+', para).group())
                text = re.sub(r'^#+\s*', '', para)
                if level == 2:
                    self.doc.add_paragraph(
                        f"{section_number}.{subsection_number} {text}",
                        style='CustomHeading2'
                    )
                    subsection_number += 1
                elif level == 3:
                    self.doc.add_paragraph(text, style='CustomHeading3')
            # Check for list items
            elif para.startswith('•') or para.startswith('-'):
                p = self.doc.add_paragraph(style='CustomBody')
                p.add_run('•').bold = True
                p.add_run(' ' + para[1:].strip())
            elif re.match(r'^\d+\.', para):
                self.doc.add_paragraph(para, style='CustomBody')
            else:
                self.doc.add_paragraph(para, style='CustomBody')

    def format_protocol(self, sections, study_type="phase1"):
        """Format complete protocol document"""
        try:
            # Get section order for study type
            section_order = self.section_orders.get(study_type, self.section_orders["phase1"])

            # Add title page
            title = self.doc.add_paragraph('Clinical Trial Protocol', style='CustomTitle')
            self.doc.add_page_break()

            # Add table of contents
            self.doc.add_paragraph("Table of Contents", style='CustomHeading1')
            self.doc.add_paragraph()  # Placeholder for TOC
            self.doc.add_page_break()

            # Process sections in order
            section_number = 1
            for section_name in section_order:
                if section_name in sections:
                    content = self._clean_section_content(sections[section_name])
                    self._format_section(section_name, content, section_number)
                    section_number += 1

            return self.doc

        except Exception as e:
            logger.error(f"Error formatting protocol: {str(e)}")
            raise

    def save_document(self, filename, format='docx'):
        """Save document in specified format"""
        try:
            if format.lower() == 'pdf':
                docx_path = f"{filename}.docx"
                pdf_path = f"{filename}.pdf"

                # Save as DOCX first
                self.doc.save(docx_path)

                # Convert to PDF
                os.system(f"libreoffice --headless --convert-to pdf {docx_path}")

                if os.path.exists(pdf_path):
                    os.remove(docx_path)
                    return pdf_path
                else:
                    raise Exception("PDF conversion failed")
            else:
                docx_path = f"{filename}.docx"
                self.doc.save(docx_path)
                return docx_path

        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            raise

def process_protocol(sections, study_type="phase1", output_format="docx"):
    """Process and format a complete protocol"""
    try:
        formatter = ProtocolFormatter()
        doc = formatter.format_protocol(sections, study_type)
        return formatter.save_document("protocol", output_format)
    except Exception as e:
        logger.error(f"Error processing protocol: {str(e)}")
        raise
