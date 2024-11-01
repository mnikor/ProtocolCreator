# protocol_post_processor.py

from bs4 import BeautifulSoup
import re
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
import os

logger = logging.getLogger(__name__)

class ProtocolPostProcessor:
    def __init__(self):
        self.section_numbers = {}
        self.current_section = None
        self.client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        self.section_order = [
            'background',
            'objectives',
            'study_design',
            'population',
            'procedures',
            'statistical',
            'safety'
        ]

    def _clean_format_with_gpt(self, content):
        """Use GPT to clean and standardize formatting"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-2024-08-06",
                messages=[
                    {
                        "role": "system",
                        "content": """Analyze and fix the protocol formatting following these rules:
1. Remove any duplicate section titles or content
2. Ensure consistent heading levels and numbering
3. Standardize bullet points and numbered lists
4. Maintain proper section order and hierarchy
5. Keep only the most detailed version of any duplicated content
6. Ensure proper spacing between sections
Do not modify any technical content or terminology."""
                    },
                    {
                        "role": "user",
                        "content": content
                    }
                ],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Error in GPT formatting: {str(e)}")
            return content

    def post_process_html(self, html_content):
        """Post-process the generated HTML content"""
        try:
            # First clean with GPT
            cleaned_content = self._clean_format_with_gpt(html_content)

            # Parse with BeautifulSoup
            soup = BeautifulSoup(cleaned_content, 'html.parser')

            # Process structural elements
            self._remove_duplicate_headings(soup)
            self._fix_heading_hierarchy(soup)
            self._fix_list_numbering(soup)
            self._fix_spacing(soup)
            self._fix_tables(soup)

            return str(soup)
        except Exception as e:
            logger.error(f"Error in HTML post-processing: {str(e)}")
            return html_content

    def _remove_duplicate_headings(self, soup):
        """Remove duplicate consecutive headings"""
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        seen_headings = set()
        for heading in headings:
            text = heading.get_text().strip().lower()
            # Remove numbering before checking for duplicates
            clean_text = re.sub(r'^\d+\.?(\d+\.?)*\s*', '', text)
            if clean_text in seen_headings:
                heading.decompose()
            else:
                seen_headings.add(clean_text)

    def _fix_heading_hierarchy(self, soup):
        """Fix heading hierarchy and add section numbering"""
        current_numbers = [0] * 6  # For h1-h6
        prev_level = 0

        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(heading.name[1]) - 1

            # Ensure proper hierarchy (no skipping levels)
            if level > prev_level + 1:
                level = prev_level + 1
                heading.name = f'h{level + 1}'

            # Reset lower level numbers
            for i in range(level + 1, 6):
                current_numbers[i] = 0

            # Increment current level
            current_numbers[level] += 1

            # Generate section number
            if level > 0:  # Don't number h1
                number = '.'.join(str(n) for n in current_numbers[:level + 1] if n > 0)
                heading.string = f"{number} {re.sub(r'^\d+\.?(\d+\.?)*\s*', '', heading.get_text().strip())}"

            prev_level = level

    def _fix_list_numbering(self, soup):
        """Fix ordered list numbering"""
        list_counter = 1
        for ol in soup.find_all('ol'):
            # Check if this is a new list or continuation
            if not ol.find_previous_sibling('ol'):
                list_counter = 1

            for li in ol.find_all('li', recursive=False):
                # Remove existing numbering
                text = li.get_text()
                text = re.sub(r'^\d+\.?\s*', '', text)
                li.string = f"{list_counter}. {text}"
                list_counter += 1

    def format_docx(self, html_content, output_path):
        """Convert processed HTML to DOCX"""
        try:
            doc = Document()

            # Set document styles
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            # Convert HTML to DOCX content
            soup = BeautifulSoup(html_content, 'html.parser')

            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'table']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    heading = doc.add_heading(element.get_text().strip(), level)
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text().strip())
                elif element.name == 'li':
                    if element.parent.name == 'ol':
                        doc.add_paragraph(element.get_text().strip(), style='List Number')
                    else:
                        doc.add_paragraph(element.get_text().strip(), style='List Bullet')
                elif element.name == 'table':
                    self._convert_table_to_docx(element, doc)

            # Add page numbers
            section = doc.sections[-1]
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.text = "Page "
            paragraph.style = doc.styles['Footer']

            doc.save(output_path)
            return True

        except Exception as e:
            logger.error(f"Error converting to DOCX: {str(e)}")
            raise

    def _convert_table_to_docx(self, html_table, doc):
        """Convert HTML table to DOCX table"""
        rows = html_table.find_all('tr')
        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                table.cell(i, j).text = cell.get_text().strip()

        doc.add_paragraph()  # Add spacing after table

def process_protocol(sections, output_format='docx'):
    """Process complete protocol with all sections"""
    try:
        processor = ProtocolPostProcessor()

        # Combine sections in proper order
        combined_html = []
        for section_name in processor.section_order:
            if section_name in sections:
                combined_html.append(sections[section_name])

        # Join and post-process content
        full_html = "\n\n".join(combined_html)
        processed_html = processor.post_process_html(full_html)

        if output_format == 'docx':
            processor.format_docx(processed_html, 'protocol.docx')
            return 'protocol.docx'
        else:
            return processed_html

    except Exception as e:
        logger.error(f"Error processing protocol: {str(e)}")
        raise